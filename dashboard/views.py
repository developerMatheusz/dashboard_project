import io
import os
from django.shortcuts import render
from django.http import JsonResponse, HttpResponse
from dotenv import load_dotenv
from docx import Document
from bs4 import BeautifulSoup
from openai import OpenAI
import psycopg2

load_dotenv()

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


def get_connection():
    conn = psycopg2.connect(
        host=os.getenv("DB_HOST"),
        port=os.getenv("DB_PORT"),
        user=os.getenv("DB_USER"),
        password=os.getenv("DB_PASSWORD"),
        dbname=os.getenv("DB_NAME")
    )
    return conn


def dashboard_page(request):
    return render(request, "dashboard.html")


def editor_page(request):
    return render(request, 'editor.html')


def xcp_data(request):
    conn = get_connection()
    cursor = conn.cursor()

    query = """
        SELECT * FROM xcp_execucao
        ORDER BY seq_execucao
    """

    cursor.execute(query)

    columns = [desc[0] for desc in cursor.description]

    data = []
    for row in cursor.fetchall():
        data.append(dict(zip(columns, row)))

    cursor.close()
    conn.close()

    return JsonResponse(data, safe=False)


def revisar_ia(request):
    if request.method == "POST":
        html_input = request.POST.get('html', '')
        soup = BeautifulSoup(html_input, 'html.parser')

        for img in soup.find_all('img'):
            img.decompose()

        texto_puro = soup.get_text(separator="\n")

        prompt = f"Você é um consultor. Melhore o texto para uma proposta comercial profissional: {texto_puro}"

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7
        )

        return JsonResponse({'revisado': response.choices[0].message.content})

    return JsonResponse({'erro': 'Método inválido'}, status=400)


def gerar_word(request):
    if request.method == "POST":
        texto_final = request.POST.get('texto_revisado', '')
        if not texto_final:
            html = request.POST.get('html_original', '')
            texto_final = BeautifulSoup(html, 'html.parser').get_text()

        doc = Document()
        doc.add_heading('Proposta Opportunya', 0)
        for linha in texto_final.split('\n'):
            if linha.strip():
                doc.add_paragraph(linha)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=proposta.docx'
        return response
